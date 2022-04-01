# author:yyr
# createDate:2021-03-29

# pip install python-docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

# 1. 创建一个文档对象
document = Document()
# Document("info.docx")  # 读取现有的word 建立文档对象

# 2. 写入内容
document.add_heading("西电电力系统公司简介", level=2)

# 样式
style = document.styles.add_style('textstyle', WD_STYLE_TYPE.PARAGRAPH)
print(style.style_id)
print(style.name)
style.font.size = Pt(5)
# 删除样式
document.styles['textstyle'].delete()

# 段落
p1 = document.add_paragraph("西安西电电力系统有限公司成立于2001年10月17日，注册地位于西安高新技术开发区西三环西辅道2号，"
                            "法定代表人为张旭宏。经营范围包括一般经营项目：灵活交流输电、高压直流输电、轻型直流输电、"
                            "电能质量、新能源及环保和节能减排工程的系统研究、成套设备设计、工程承包、设备制造、销售及技术咨询服务；"
                            "电力电子产品研发、设计、制造、销售；固定串补/可控串补设备、无功补偿和谐波治理设备、可控高抗设备、"
                            "直流输电换流阀、电力控制保护设备、直流场设备、储能装置设备、新能源变流器、牵引变流器、特种电源、"
                            "中高压变频器、高压软启动产品、电力电子器件的销售；", style="textstyle")
p1.insert_paragraph_before("!!!在段落前插入一个新的段落")
format = p1.paragraph_format
# 左右缩进
format.left_indent = Pt(20)
format.right_indent = Pt(20)
# 首行缩进
format.first_line_indent = Pt(20)
# 行间距
format.line_spacing = 1.5
# 追加
run = p1.add_run("西安西电电力系统有限公司对外投资4家公司。")
# 字体、字号、文字颜色
run.font.size = Pt(12)
run.font.name = "微软雅黑"
run.font.color.rgb = RGBColor(235, 33, 24)
run1 = p1.add_run("包括：西安城投西电智能充电有限公司、西安端怡科技有限公司、西电集团财务有限责任公司、陕西半导体先导技术中心有限公司等。")
# 加粗、下划线、斜体
run1.bold = True
run1.font.underline = True
run1.font.italic = True

# 插入图片
document.add_picture("gmt.png")
document.add_picture("gmt.png", Pt(30), Pt(20))

#  插入表格
table = document.add_table(rows=1, cols=3, style="Medium List 2")
# 表头
header_cells = table.rows[0].cells
header_cells[0].text = "月份"
header_cells[1].text = "预期销售额"
header_cells[2].text = "实际销售额"
# 数据
data = (
    ['一月份', 300, 600],
    ['二月份', 400, 700],
    ['三月份', 500, 800],
)
for item in data:
    rows_cells = table.add_row().cells
    rows_cells[0].text = item[0]
    rows_cells[1].text = str(item[1])
    rows_cells[2].text = str(item[2])

# 获取表格
print(len(document.tables[0].rows))  # 打印总行数
print(len(document.tables[0].columns))  # 打印总列数
# cell
print(document.tables[0].cell(0, 2).text)

# 3. 保存文档
document.save("info.docx")
